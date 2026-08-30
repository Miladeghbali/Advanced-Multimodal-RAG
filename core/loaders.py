from __future__ import annotations

import csv
from pathlib import Path
from typing import Iterable

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".csv", ".html", ".htm", ".docx"}


def _base_metadata(path: Path, page: int = 1) -> dict:
    return {
        "source": path.name,
        "page": page,
        "source_type": path.suffix.lower().lstrip("."),
    }


def _load_pdf(path: Path) -> list[Document]:
    reader = PdfReader(str(path), strict=False)
    docs: list[Document] = []
    for page_index, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(
                Document(
                    page_content=text,
                    metadata=_base_metadata(path, page_index + 1),
                )
            )
    return docs


def _load_text(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [Document(page_content=text, metadata=_base_metadata(path))]


def _load_csv(path: Path) -> list[Document]:
    rows: list[str] = []
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as fh:
        reader = csv.DictReader(fh)
        if reader.fieldnames:
            for row_index, row in enumerate(reader, start=1):
                row_text = " | ".join(f"{key}: {value}" for key, value in row.items())
                rows.append(f"row {row_index}: {row_text}")
        else:
            fh.seek(0)
            plain_reader = csv.reader(fh)
            rows = [" | ".join(row) for row in plain_reader]
    return [Document(page_content="\n".join(rows), metadata=_base_metadata(path))]


def _load_html(path: Path) -> list[Document]:
    html = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = "\n".join(line.strip() for line in soup.get_text("\n").splitlines() if line.strip())
    return [Document(page_content=text, metadata=_base_metadata(path))]


def _load_docx(path: Path) -> list[Document]:
    docx = DocxDocument(str(path))
    text = "\n".join(p.text for p in docx.paragraphs if p.text.strip())
    return [Document(page_content=text, metadata=_base_metadata(path))]


def load_file(path: str | Path) -> list[Document]:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix}")
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".csv":
        return _load_csv(path)
    if suffix in {".html", ".htm"}:
        return _load_html(path)
    if suffix == ".docx":
        return _load_docx(path)
    return _load_text(path)


def load_files(paths: Iterable[str | Path]) -> list[Document]:
    docs: list[Document] = []
    for path in paths:
        docs.extend(load_file(path))
    return docs
